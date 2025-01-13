import os
import json
import docx
import pypdf
import pandas as pd


class FileLoader:
    def __init__(self, allowed_formats=None):
        """
        Инициализация класса FileLoader

        :param allowed_formats: Список допустимых форматов файлов
        """
        self.files = []
        self.allowed_formats = allowed_formats or [
            '.txt', '.pdf', '.docx', '.csv', '.xlsx'
        ]

    def get_allowed_formats(self):
        """
        Получение списка доступных форматов файлов

        :return: Список разрешенных форматов файлов
        """
        return self.allowed_formats.copy()

    def add_allowed_format(self, file_format):
        """
        Добавление нового формата файла в список разрешенных

        :param file_format: Формат файла (например, '.json')
        :return: True, если формат добавлен, False если уже существует
        """
        # Нормализация формата (добавление точки если отсутствует)
        if not file_format.startswith('.'):
            file_format = '.' + file_format

        if file_format not in self.allowed_formats:
            self.allowed_formats.append(file_format)
            return True
        return False

    def remove_allowed_format(self, file_format):
        """
        Удаление формата из списка разрешенных

        :param file_format: Формат файла для удаления
        :return: True, если формат удален, False если не найден
        """
        # Нормализация формата (добавление точки если отсутствует)
        if not file_format.startswith('.'):
            file_format = '.' + file_format

        if file_format in self.allowed_formats:
            # Нельзя удалить все форматы
            if len(self.allowed_formats) > 1:
                self.allowed_formats.remove(file_format)
                return True
            else:
                print("Нельзя удалить последний формат файла")
                return False
        return False

    def get_file_list(self, filter_format=None):
        """
        Получение списка файлов с возможностью фильтрации по формату

        :param filter_format: Формат файла для фильтрации (необязательно)
        :return: Список файлов
        """
        # Если формат не указан, возвращаем все файлы
        if not filter_format:
            return self.files.copy()

        # Нормализация формата
        if not filter_format.startswith('.'):
            filter_format = '.' + filter_format

        # Фильтрация файлов по указанному формату
        return [
            file for file in self.files
            if os.path.splitext(file)[1].lower() == filter_format
        ]

    def add_file(self, file_path):
        """
        Добавление отдельного файла в список

        :param file_path: Путь к файлу
        :return: True, если файл добавлен успешно, иначе False
        """
        # Проверка расширения файла
        file_ext = os.path.splitext(file_path)[1].lower()
        if file_ext not in self.allowed_formats:
            print(f"Формат файла {file_ext} не поддерживается")
            return False

        if not os.path.exists(file_path):
            print("Файл не существует")
            return False

        # Проверка на дублирование
        if file_path in self.files:
            print("Файл уже добавлен")
            return False

        # Добавление файла
        self.files.append(file_path)
        return True

    def add_files_from_folder(self, folder_path):
        """
        Добавление всех файлов из указанной папки

        :param folder_path: Путь к папке
        :return: Количество добавленных файлов
        """
        if not os.path.exists(folder_path):
            print(f"Папка '{folder_path}' не существует")
            return 0

        if not os.path.isdir(folder_path):
            print(f"Путь '{folder_path}' не является папкой")
            return 0

        added_count = 0
        for filename in os.listdir(folder_path):
            file_path = os.path.join(folder_path, filename)
            if os.path.isfile(file_path):
                if self.add_file(file_path):
                    added_count += 1
        return added_count

    def remove_file(self, file_path):
        """
        Удаление файла из списка

        :param file_path: Путь к файлу для удаления
        :return: True, если файл удален, иначе False
        """
        if file_path not in self.files:
            print(f"Файл '{file_path}' отсутствует в списке")
            return False

        self.files.remove(file_path)
        print(f"Файл '{file_path}' удален из списка")
        return True

    def remove_files_from_folder(self, folder_path):
        """
        Удаление файлов, добавленных из указанной папки

        :param folder_path: Путь к папке
        :return: Количество удаленных файлов
        """
        # Нормализуем путь к папке для точного сравнения
        folder_path = os.path.abspath(folder_path)

        removed_count = 0
        # Создаем новый список файлов, исключая те, которые находятся в указанной папке
        new_files = []
        for file_path in self.files:
            if os.path.abspath(os.path.dirname(file_path)) == folder_path:
                removed_count += 1
            else:
                new_files.append(file_path)

        # Обновляем список файлов
        self.files = new_files
        return removed_count

    def clear_file_list(self):
        """
        Очистка списка файлов

        :return: True, если список очищен
        """
        self.files.clear()
        print("Список файлов успешно очищен.")
        return True

    def get_file_content(self, file_path, exclude_types=None):
        """
        Получение содержимого файла с возможностью исключения определенных элементов

        :param file_path: Путь к файлу
        :param exclude_types: Список типов контента для исключения
        :return: Текстовое содержимое файла
        """
        if file_path not in self.files:
            print("Файл не в списке")
            return None

        exclude_types = exclude_types or []
        file_ext = os.path.splitext(file_path)[1].lower()

        try:
            if file_ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()

            elif file_ext == '.pdf':
                with open(file_path, 'rb') as f:
                    reader = pypdf.PdfReader(f)
                    text = " ".join([page.extract_text()
                                    for page in reader.pages])
                    return text

            elif file_ext == '.docx':
                doc = docx.Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])

                if 'tables' in exclude_types:
                    text = "\n".join(
                        [para.text for para in doc.paragraphs if not para.text.startswith('|')])

                return text

            elif file_ext in ['.csv', '.xlsx']:
                if 'tables' in exclude_types:
                    return None

                if file_ext == '.csv':
                    df = pd.read_csv(file_path)
                else:
                    df = pd.read_excel(file_path)

                return df.to_string()

        except Exception as e:
            print(f"Ошибка чтения файла: {e}")
            return None

    def save_file_list(self, output_path='file_list.json'):
        """
        Сохранение списка файлов в JSON

        :param output_path: Путь для сохранения списка
        """
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(self.files, f)

    def load_file_list(self, input_path='file_list.json'):
        """
        Загрузка списка файлов из JSON

        :param input_path: Путь к файлу со списком
        """
        try:
            with open(input_path, 'r', encoding='utf-8') as f:
                loaded_files = json.load(f)
                # Проверяем существование каждого файла
                self.files = [f for f in loaded_files if os.path.exists(f)]
        except FileNotFoundError:
            print("Файл списка не найден")
